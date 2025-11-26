"""Tests for DOCX document handler."""

import tempfile
from pathlib import Path

from docx import Document

from anonymizer.handlers.docx_handler import DocxHandler


class TestDocxHandler:
    """Tests for DocxHandler class."""

    def test_supported_extensions(self) -> None:
        """Test that handler reports correct extensions."""
        handler = DocxHandler()
        assert handler.supported_extensions == (".docx",)

    def test_read_file(self) -> None:
        """Test reading a docx file."""
        handler = DocxHandler()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "test.docx"

            doc = Document()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("Line 2")
            doc.save(str(path))

            content = handler.read(path)

            assert "Hello World" in content
            assert "Line 2" in content

    def test_write_file(self) -> None:
        """Test writing a docx file."""
        handler = DocxHandler()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "output.docx"

            handler.write(path, "Test content\nLine 2")

            assert path.exists()

            doc = Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs]
            assert "Test content" in paragraphs
            assert "Line 2" in paragraphs

    def test_write_creates_parent_dirs(self) -> None:
        """Test that write creates parent directories."""
        handler = DocxHandler()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "nested" / "dir" / "output.docx"

            handler.write(path, "Content")

            assert path.exists()

    def test_read_with_table(self) -> None:
        """Test reading docx file with tables."""
        handler = DocxHandler()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "test_table.docx"

            doc = Document()
            doc.add_paragraph("Before table")
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Cell 1"
            table.cell(0, 1).text = "Cell 2"
            table.cell(1, 0).text = "Cell 3"
            table.cell(1, 1).text = "Cell 4"
            doc.add_paragraph("After table")
            doc.save(str(path))

            content = handler.read(path)

            assert "Before table" in content
            assert "After table" in content
            assert "Cell 1" in content
            assert "Cell 4" in content

    def test_read_unicode(self) -> None:
        """Test reading file with unicode characters."""
        handler = DocxHandler()

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "unicode.docx"

            doc = Document()
            doc.add_paragraph("María García está en España")
            doc.save(str(path))

            content = handler.read(path)

            assert "María García" in content
            assert "España" in content
