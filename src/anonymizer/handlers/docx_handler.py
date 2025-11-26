"""Handler for Word documents (.docx)."""

from pathlib import Path
from typing import List, Tuple

from docx import Document  # type: ignore[import-untyped]
from docx.table import Table  # type: ignore[import-untyped]
from docx.text.paragraph import Paragraph  # type: ignore[import-untyped]

from ..logger import setup_logger

logger = setup_logger(__name__)


class DocxHandler:
    """
    Handler for Word (.docx) documents.

    Extracts text from paragraphs and tables while preserving structure.
    """

    @property
    def supported_extensions(self) -> Tuple[str, ...]:
        """Return supported file extensions."""
        return (".docx",)

    def read(self, path: Path) -> str:
        """
        Extract text from a Word document.

        Extracts text from paragraphs, tables, headers, and footers.

        Args:
            path: Path to the .docx file

        Returns:
            Extracted text content with preserved structure
        """
        logger.info(f"reading docx file path:{path}")

        doc = Document(path)
        text_parts: List[str] = []

        text_parts.extend(self._extract_paragraphs(doc))
        text_parts.extend(self._extract_tables(doc))
        text_parts.extend(self._extract_headers_footers(doc))

        content = "\n".join(text_parts)

        logger.info(f"docx file read path:{path};length:{len(content)}")
        return content

    def _extract_paragraphs(self, doc: Document) -> List[str]:
        """Extract text from all paragraphs."""
        return [para.text for para in doc.paragraphs if para.text.strip()]

    def _extract_tables(self, doc: Document) -> List[str]:
        """Extract text from all tables."""
        table_texts: List[str] = []

        for table in doc.tables:
            table_texts.extend(self._extract_table_text(table))

        return table_texts

    def _extract_table_text(self, table: Table) -> List[str]:
        """Extract text from a single table."""
        texts: List[str] = []

        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                texts.append(" | ".join(row_texts))

        return texts

    def _extract_headers_footers(self, doc: Document) -> List[str]:
        """Extract text from headers and footers."""
        texts: List[str] = []

        for section in doc.sections:
            header_text = self._get_header_footer_text(section.header)
            footer_text = self._get_header_footer_text(section.footer)

            if header_text:
                texts.append(header_text)
            if footer_text:
                texts.append(footer_text)

        return texts

    def _get_header_footer_text(self, header_footer: Paragraph) -> str:
        """Extract text from a header or footer."""
        paragraphs = [p.text for p in header_footer.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)

    def write(self, path: Path, content: str) -> None:
        """
        Write anonymized content to a Word document.

        Creates a new document with the anonymized text.

        Args:
            path: Path where to save the .docx file
            content: Anonymized text content
        """
        logger.info(f"writing docx file path:{path}")

        path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        paragraphs = content.split("\n")

        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)

        doc.save(path)

        logger.info(f"docx file written path:{path};paragraphs:{len(paragraphs)}")
